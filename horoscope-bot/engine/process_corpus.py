#!/usr/bin/env python3
"""
Batch corpus processor — handles ALL remaining documents.

Usage:
    python3 engine/process_corpus.py        # process everything
    python3 engine/process_corpus.py --only giao_trinh,client,choudhry
    python3 engine/process_corpus.py --only bphs,nakshatra,bhrigu
"""

import json, hashlib, re, sys
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent
RAW_DIR = BASE_DIR / "corpus" / "raw"
CHUNKS_DIR = BASE_DIR / "corpus" / "chunks"
CHUNKS_DIR.mkdir(parents=True, exist_ok=True)


def clean_text(text):
    """Normalize whitespace."""
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def make_chunk_id(prefix, text):
    cid = hashlib.md5(text[:80].encode()).hexdigest()[:12]
    return f"{prefix}_{cid}"


def save_chunks(chunks, name):
    path = CHUNKS_DIR / f"{name}.jsonl"
    count_before = 0
    if path.exists():
        count_before = sum(1 for _ in open(path))
    with open(path, 'w', encoding='utf-8') as f:
        for chunk in chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + '\n')
    count_after = len(chunks)
    print(f"  💾 {name}: {count_before} → {count_after} chunks ({count_after-count_before:+d})")
    return path


# ──────────────────────────────────────────────
# 1. GIÁO TRÌNH TIẾNG VIỆT
# ──────────────────────────────────────────────

def process_giao_trinh():
    print("\n📖 [GIÁO TRÌNH TIẾNG VIỆT]")
    chunks = []
    
    # 12 cung Hoàng Đạo
    src = RAW_DIR / "Tìm hiểu về 12 cung Hoàng Đạo.docx"
    if src.exists():
        from docx import Document
        doc = Document(str(src))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Split by cung names
        current_cung = "Mở đầu"
        current_text = []
        for p in paras:
            # Detect cung name
            match = re.match(r'^(Cung\s+\S+|Bài\s+\d+|I{1,3}V?|V?I{0,3})\b', p)
            if match and len(current_text) > 0:
                if current_text:
                    chunk_text = '\n'.join(current_text)
                    chunks.append({
                        "id": make_chunk_id("gt_cung", chunk_text),
                        "text": chunk_text,
                        "metadata": {
                            "type": "giao_trinh",
                            "source": "12_cung_hoang_dao",
                            "topic": current_cung,
                            "lang": "vi",
                            "char_count": len(chunk_text),
                        }
                    })
                current_cung = p[:50]
                current_text = [p]
            else:
                current_text.append(p)
        
        # Last chunk
        if current_text:
            chunk_text = '\n'.join(current_text)
            chunks.append({
                "id": make_chunk_id("gt_cung", chunk_text),
                "text": chunk_text,
                "metadata": {
                    "type": "giao_trinh",
                    "source": "12_cung_hoang_dao",
                    "topic": current_cung,
                    "lang": "vi",
                    "char_count": len(chunk_text),
                }
            })
        print(f"  📄 12 cung Hoàng Đạo: {len(chunks)} chunks")
    
    # Giáo án hành tinh
    src2 = RAW_DIR / "giao an hanh tinh.docx"
    if src2.exists():
        from docx import Document
        doc = Document(str(src2))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Split by hành tinh name
        hanh_tinh_names = ["Mặt Trời", "Mặt Trăng", "Sao Thủy", "Sao Kim", "Sao Hỏa", 
                           "Sao Mộc", "Sao Thổ", "Sao Thiên Vương", "Sao Hải Vương", "Sao Diêm Vương"]
        
        current_planet = "Mở đầu"
        current_text = []
        planet_chunks = 0
        
        for p in paras:
            matched = False
            for ht in hanh_tinh_names:
                if p.startswith(ht) or p.startswith(ht.upper()):
                    if current_text:
                        chunk_text = '\n'.join(current_text)
                        chunks.append({
                            "id": make_chunk_id("gt_hanhtinh", chunk_text),
                            "text": chunk_text,
                            "metadata": {
                                "type": "giao_trinh",
                                "source": "giao_an_hanh_tinh",
                                "planet": current_planet,
                                "lang": "vi",
                                "char_count": len(chunk_text),
                            }
                        })
                        planet_chunks += 1
                    current_planet = p[:50]
                    current_text = [p]
                    matched = True
                    break
            if not matched:
                current_text.append(p)
        
        if current_text:
            chunk_text = '\n'.join(current_text)
            chunks.append({
                "id": make_chunk_id("gt_hanhtinh", chunk_text),
                "text": chunk_text,
                "metadata": {
                    "type": "giao_trinh",
                    "source": "giao_an_hanh_tinh",
                    "planet": current_planet,
                    "lang": "vi",
                    "char_count": len(chunk_text),
                }
            })
            planet_chunks += 1
        
        print(f"  📄 Giáo án hành tinh: {planet_chunks} planet chunks")
    
    if chunks:
        save_chunks(chunks, "giao_trinh_viet")
    return len(chunks)


# ──────────────────────────────────────────────
# 2. CLIENT READINGS (few-shot examples)
# ──────────────────────────────────────────────

def process_client_readings():
    print("\n📖 [CLIENT READINGS — FEW-SHOT]")
    chunks = []
    
    clients = [
        ("Nguyễn Ngọc Quỳnh-24_12_1991(3h51p).docx", "client_quynh"),
        ("Trà Mình Trí .docx", "client_tra_minh_tri"),
    ]
    
    from docx import Document
    
    for fname, cid in clients:
        src = RAW_DIR / fname
        if not src.exists():
            continue
        doc = Document(str(src))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = '\n'.join(paras)
        
        # Split into meaningful chunks (by headings or sections)
        sections = re.split(r'\n(?=[A-Z][A-Za-z\s]+:|\d+\.\s+[A-Z])', full_text)
        
        for sec in sections:
            sec = clean_text(sec)
            if len(sec) < 100:
                continue
            heading = sec.split('\n')[0][:60]
            chunks.append({
                "id": make_chunk_id(cid, sec),
                "text": sec,
                "metadata": {
                    "type": "client_reading",
                    "client": cid,
                    "heading": heading,
                    "lang": "vi",
                    "char_count": len(sec),
                }
            })
        
        print(f"  📄 {fname}: {len(sections)} sections → {sum(1 for s in sections if len(clean_text(s))>=100)} chunks")
    
    # PDF client readings
    pdf_clients = [
        ("ho_thi_linh.txt", "client_ho_thi_linh"),
        ("quynh_vy.txt", "client_quynh_vy"),
    ]
    
    for fname, cid in pdf_clients:
        src = RAW_DIR / fname
        if not src.exists():
            continue
        text = src.read_text(encoding='utf-8')
        sections = re.split(r'\n(?=[A-Z][A-Za-z\s]+:|\d+\.\s+[A-Z]|\bPHẦN\s+\d+)', text)
        
        for sec in sections:
            sec = clean_text(sec)
            if len(sec) < 100:
                continue
            heading = sec.split('\n')[0][:60]
            chunks.append({
                "id": make_chunk_id(cid, sec),
                "text": sec,
                "metadata": {
                    "type": "client_reading",
                    "client": cid,
                    "heading": heading,
                    "lang": "vi",
                    "char_count": len(sec),
                }
            })
        
        print(f"  📄 {fname}: {len(chunks)} chunks so far")
    
    if chunks:
        save_chunks(chunks, "client_readings")
    return len(chunks)


# ──────────────────────────────────────────────
# 3. CHOUDHRY — LAGNA SIGNS
# ──────────────────────────────────────────────

def process_choudhry():
    print("\n📖 [CHOUDHRY — LAGNA SIGNS]")
    src = RAW_DIR / "choudhry_lagna_signs.txt"
    if not src.exists():
        print("  ❌ File not found")
        return 0
    
    text = src.read_text(encoding='utf-8')
    
    signs = ["Aries", "Taurus", "Gemini", "Cancer", "Leo", "Virgo",
             "Libra", "Scorpio", "Sagittarius", "Capricorn", "Aquarius", "Pisces"]
    
    chunks = []
    current_sign = "Mở đầu"
    current_text = []
    
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        
        # Detect sign heading
        matched = False
        for s in signs:
            # Pattern: "Aries" or "Chapter X: Aries"
            if re.match(rf'^(Chapter\s+\d+\s*[:.-]?\s*)?{re.escape(s)}\b', stripped, re.IGNORECASE):
                if current_text:
                    chunk_text = '\n'.join(current_text)
                    chunks.append({
                        "id": make_chunk_id("choudhry", chunk_text),
                        "text": chunk_text,
                        "metadata": {
                            "type": "lagna_sign",
                            "source": "choudhry",
                            "sign": current_sign,
                            "char_count": len(chunk_text),
                        }
                    })
                current_sign = s
                current_text = [stripped]
                matched = True
                break
        
        if not matched:
            current_text.append(stripped)
    
    # Last sign
    if current_text:
        chunk_text = '\n'.join(current_text)
        chunks.append({
            "id": make_chunk_id("choudhry", chunk_text),
            "text": chunk_text,
            "metadata": {
                "type": "lagna_sign",
                "source": "choudhry",
                "sign": current_sign,
                "char_count": len(chunk_text),
            }
        })
    
    print(f"  📄 {len(chunks)} sign chunks from Choudhry")
    if chunks:
        save_chunks(chunks, "choudhry_lagna")
    return len(chunks)


# ──────────────────────────────────────────────
# 4. FIX 6 MISSING NAKSHATRAS
# ──────────────────────────────────────────────

def fix_missing_nakshatras():
    print("\n📖 [NAKSHATRAS — 6 MISSING]")
    src = RAW_DIR / "book-of-naksatras.txt"
    if not src.exists():
        print("  ❌ File not found")
        return 0
    
    text = src.read_text(encoding='utf-8')
    pages = text.split('\n---PAGE BREAK---\n')
    
    # The 6 missing: Purva Phalguni, Jyeshtha, Purva Ashadha, 
    # Uttara Ashadha, Purva Bhadrapada, Uttara Bhadrapada
    missing = [
        ("Purva Phalguni", ["Purva Phalguni", "Purva Phalguni"]),
        ("Jyeshtha", ["Jyeshtha"]),
        ("Purva Ashadha", ["Purva Ashadha", "Purva Ashada"]),
        ("Uttara Ashadha", ["Uttara Ashadha", "Uttara Ashada"]),
        ("Purva Bhadrapada", ["Purva Bhadrapada", "Purva Bhadrapada"]),
        ("Uttara Bhadrapada", ["Uttara Bhadrapada", "Uttara Bhadrapada"]),
    ]
    
    # Find pages containing these nakshatras
    found = {nk: [] for nk, _ in missing}
    
    for pg_idx, page_text in enumerate(pages):
        first_300 = page_text[:300].lower()
        for nk_name, variants in missing:
            for v in variants:
                if v.lower() in first_300:
                    found[nk_name].append(pg_idx)
                    break
    
    # Load existing nakshatra chunks to avoid duplicates
    existing_path = CHUNKS_DIR / "nakshatras.jsonl"
    existing_naks = set()
    if existing_path.exists():
        with open(existing_path) as f:
            for line in f:
                chunk = json.loads(line)
                existing_naks.add(chunk["metadata"].get("name", ""))
    
    print(f"  📊 Already have: {sorted(existing_naks)}")
    
    new_chunks = []
    for nk_name, variants in missing:
        if nk_name in existing_naks:
            print(f"  ⏭️  {nk_name} — already exists")
            continue
        
        pages_found = found.get(nk_name, [])
        if not pages_found:
            print(f"  ⚠️  {nk_name} — not found in pages")
            # Try searching full text
            for pg_idx, page_text in enumerate(pages):
                if nk_name.lower() in page_text[:300].lower():
                    pages_found.append(pg_idx)
            if not pages_found:
                print(f"  ❌ {nk_name} — really not found")
                continue
        
        # Get text from found pages
        start_pg = min(pages_found)
        end_pg = max(pages_found) + 2  # next 2 pages for safety
        
        section_text = '\n'.join(pages[start_pg:end_pg])
        section_text = clean_text(section_text)
        
        new_chunks.append({
            "id": f"nak_{nk_name.lower().replace(' ','_')}_{make_chunk_id('fix', section_text)}",
            "text": section_text,
            "metadata": {
                "type": "nakshatra",
                "name": nk_name,
                "source": "book-of-naksatras",
                "pages": f"{start_pg+1}-{end_pg}",
                "char_count": len(section_text),
            }
        })
        print(f"  ✅ {nk_name}: page {start_pg+1}-{end_pg}, {len(section_text)} chars")
    
    if new_chunks:
        # Append to existing nakshatras.jsonl
        with open(existing_path, 'a', encoding='utf-8') as f:
            for chunk in new_chunks:
                f.write(json.dumps(chunk, ensure_ascii=False) + '\n')
        print(f"  💾 Appended {len(new_chunks)} chunks to nakshatras.jsonl")
    
    return len(new_chunks)


# ──────────────────────────────────────────────
# 5. BPHS — FIX CHAPTER DETECTION
# ──────────────────────────────────────────────

def process_bphs():
    print("\n📖 [BPHS VOL 1 — CHAPTER FIX]")
    
    # Try English OCR first
    src = RAW_DIR / "bphs_ocr_english.txt"
    if src.exists():
        text = src.read_text(encoding='utf-8')
    else:
        print("  ❌ No OCR text found")
        return 0
    
    # BPHS chapters: detect patterns like "Chapter 1", "CH. 1", "1."
    # Also specific chapter names
    chapters = re.split(
        r'(?:^|\n)(?:Chapter\s+\d+|CH\.\s+\d+|CHAPTER\s+\d+)\b[^\n]*',
        text, flags=re.IGNORECASE
    )
    
    if len(chapters) < 3:
        # Try alternative: detect by numbered sections
        chapters = re.split(r'\n(?:\d{1,2}\.\s+[A-Z][^\n]+)', text)
    
    if len(chapters) < 3:
        # Just split into reasonably-sized chunks
        print("  ⚠️ Chapter detection limited, using paragraph chunks")
        paras = text.split('\n')
        chapters = []
        current = []
        for p in paras:
            if len(current) > 20:  # ~20 lines per chunk
                chapters.append('\n'.join(current))
                current = []
            current.append(p)
        if current:
            chapters.append('\n'.join(current))
    
    chunks = []
    for i, chap in enumerate(chapters):
        chap = clean_text(chap)
        if len(chap) < 200:
            continue
        
        heading = chap.split('\n')[0][:80] if chap.split('\n')[0] else f"Section {i+1}"
        
        chunks.append({
            "id": make_chunk_id("bphs", chap),
            "text": chap,
            "metadata": {
                "type": "bphs_chapter",
                "chapter": i + 1,
                "heading": heading,
                "source": "BPHS_Vol1_OCR",
                "char_count": len(chap),
            }
        })
    
    print(f"  📄 {len(chunks)} chapter chunks from BPHS")
    if chunks:
        save_chunks(chunks, "bphs_vol1")
    return len(chunks)


# ──────────────────────────────────────────────
# 6. BHRIGU SAMHITA — SAMPLE
# ──────────────────────────────────────────────

def process_bhrigu():
    print("\n📖 [BHRIGU SAMHITA]")
    src = RAW_DIR / "bhrigu_samhita.txt"
    if not src.exists():
        print("  ❌ File not found")
        return 0
    
    text = src.read_text(encoding='utf-8')
    
    # Bhrigu Samhita is massive (1.7M chars). 
    # Split into chapters or major sections
    # It's structured with numbered slokas/sections
    
    # Try detecting chapter/section boundaries
    sections = re.split(r'\n(?=(?:CHAPTER|SECTION|SLOKA|Chapter|Section|Sloka)\s+\d+)', text)
    
    if len(sections) < 5:
        # Fallback: split by page-like boundaries
        sections = re.split(r'\n(?=\d+\s*\n)', text)
    
    if len(sections) < 5:
        # Last resort: fixed-size chunks
        paras = text.split('\n')
        sections = []
        for i in range(0, len(paras), 30):
            sec = '\n'.join(paras[i:i+30])
            if len(clean_text(sec)) > 100:
                sections.append(sec)
    
    chunks = []
    for i, sec in enumerate(sections):
        sec = clean_text(sec)
        if len(sec) < 200:
            continue
        
        heading = sec.split('\n')[0][:80] if sec.split('\n')[0] else f"Section {i+1}"
        
        chunks.append({
            "id": make_chunk_id("bhrigu", sec),
            "text": sec,
            "metadata": {
                "type": "bhrigu_samhita",
                "section": i + 1,
                "heading": heading,
                "source": "Bhrigu_Samhita",
                "char_count": len(sec),
            }
        })
        
        # Limit to 200 chunks max (Bhrigu is huge)
        if len(chunks) >= 200:
            break
    
    print(f"  📄 {len(chunks)} sections from Bhrigu Samhita")
    if chunks:
        save_chunks(chunks, "bhrigu_samhita")
    return len(chunks)


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────

def run(only=None):
    print("=" * 60)
    print("🔮 VOTIVE ACADEMY — Batch Corpus Processor")
    print("=" * 60)
    
    tasks = {
        "giao_trinh": process_giao_trinh,
        "client": process_client_readings,
        "choudhry": process_choudhry,
        "nakshatra": fix_missing_nakshatras,
        "bphs": process_bphs,
        "bhrigu": process_bhrigu,
    }
    
    if only:
        selected = [t.strip() for t in only.split(',')]
        tasks_to_run = {k: v for k, v in tasks.items() if k in selected}
    else:
        tasks_to_run = tasks
    
    total = 0
    for name, fn in tasks_to_run.items():
        count = fn()
        total += count
    
    print(f"\n{'='*60}")
    print(f"✅ TOTAL: {total} new chunks processed")
    
    # Final count
    all_chunks = 0
    for f in sorted(CHUNKS_DIR.glob("*.jsonl")):
        count = sum(1 for _ in open(f))
        all_chunks += count
        print(f"  {f.name}: {count}")
    print(f"📊 FINAL: {all_chunks} total chunks")
    print(f"{'='*60}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--only", type=str, help="Comma-separated: giao_trinh,client,choudhry,nakshatra,bphs,bhrigu")
    args = parser.parse_args()
    run(args.only)
